# W.I.T.N.E.S.S. - Web-based Interrogation and Testimony via a Neural Engaged Speech System
# Copyright (C) 2026 Philip Roy <https://www.bluengrey.com>
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program. If not, see <https://www.gnu.org/licenses/>.

import sys
import os
import base64
from docx.shared import Inches
from backend.models.audio.audio_handler import transcribe_audio_file, synthesize_speech
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import re

def _two_digit_words(n: int) -> str:
    ones = ["zero","one","two","three","four","five","six","seven","eight","nine"]
    teens = ["ten","eleven","twelve","thirteen","fourteen","fifteen","sixteen","seventeen","eighteen","nineteen"]
    tens = ["","","twenty","thirty","forty","fifty","sixty","seventy","eighty","ninety"]
    if n < 10:
        return ones[n]
    if 10 <= n < 20:
        return teens[n-10]
    t, o = divmod(n, 10)
    return tens[t] + ("-" + ones[o] if o else "")

# Digit-to-word mapping for phone number pronunciation
_DIGIT_WORDS = {
    "0": "zero", "1": "one", "2": "two", "3": "three", "4": "four",
    "5": "five", "6": "six", "7": "seven", "8": "eight", "9": "nine"
}

def _phone_to_tts(match: re.Match) -> str:
    """Convert phone number match to TTS-friendly format with individual digits.

    Examples:
        "021-123-4567" -> "zero two one, one two three, four five six seven"
        "09-555-1234" -> "zero nine, five five five, one two three four"
    """
    phone = match.group(0)
    # Split by common phone separators (hyphens, spaces, parentheses, dots)
    groups = re.split(r"[-\s().]+", phone.strip())
    spoken_groups = []
    for group in groups:
        if not group:
            continue
        # Convert each digit to word
        spoken = " ".join(_DIGIT_WORDS.get(d, d) for d in group)
        spoken_groups.append(spoken)
    # Join groups with comma for natural pause
    return ", ".join(spoken_groups)


def normalize_tts_nz(text: str) -> str:
    """Prepare text for Piper TTS so it pronounces NZ-specific content correctly.

    This function converts written symbols to spoken forms before the text
    reaches the TTS engine. Applied to LLM output before synthesis.
    """
    # NZ emergency number — must be spoken as three separate digits, not "one hundred and eleven"
    text = re.sub(r"\b111\b", "one one one", text)

    # Years: years are spoken in two-digit pairs in NZ English.
    # "1993" -> "nineteen ninety-three", not "one thousand nine hundred and ninety-three"
    # "2004" -> "two thousand and four" (special case: 2000–2009 can't use the pair pattern)
    def repl_year_19(m):
        return "nineteen " + _two_digit_words(int(m.group(1)))
    def repl_year_20(m):
        last_two = m.group(1)
        num = int(last_two)
        if num == 0:
            return "two thousand"
        elif num < 10:
            return "two thousand and " + _two_digit_words(num)
        else:
            return "twenty " + _two_digit_words(num)
    text = re.sub(r"\b19(\d{2})\b", repl_year_19, text)
    text = re.sub(r"\b20(\d{2})\b", repl_year_20, text)

    # Phone numbers must be spoken digit-by-digit in NZ English (e.g., "zero two one, one two three...").
    # This must run BEFORE the number-ranges substitution below, otherwise "021-123" would become "021 to 123".
    text = re.sub(r"\(?\d{2,3}\)?[-\s.]\d{3,4}[-\s.]\d{3,4}\b", _phone_to_tts, text)

    # Number ranges like "6-7 metres" should be spoken "six to seven metres", not "six dash seven metres"
    text = re.sub(r"\b(\d+)-(\d+)\b", r"\1 to \2", text)

    return text

# --- NZ English spelling normaliser for on-screen text (LLM + transcription) ---

def _preserve_case(src: str, repl: str) -> str:
    """Return `repl` with the casing pattern of `src` (UPPER, Title, lower)."""
    if src.isupper():
        return repl.upper()
    if src[:1].isupper() and src[1:].islower():
        return repl.capitalize()
    return repl

# Word-boundary, case-insensitive substitution pairs.
# Applied by normalize_en_nz() to convert American spellings to NZ English
# in LLM output and STT transcripts before they reach the user interface.
_NZ_SPELLING_RULES = [
    (r"\bfavor\b", "favour"), (r"\bfavors\b", "favours"), (r"\bfavored\b", "favoured"), (r"\bfavoring\b", "favouring"),
    (r"\bfavorite\b", "favourite"), (r"\bfavorites\b", "favourites"),
    (r"\bcolor\b", "colour"), (r"\bcolors\b", "colours"), (r"\bcolored\b", "coloured"), (r"\bcoloring\b", "colouring"),
    (r"\bhonor\b", "honour"), (r"\bhonors\b", "honours"), (r"\bhonored\b", "honoured"), (r"\bhonoring\b", "honouring"),
    (r"\bbehavior\b", "behaviour"), (r"\bbehaviors\b", "behaviours"),
    (r"\bneighbor\b", "neighbour"), (r"\bneighbors\b", "neighbours"), (r"\bneighboring\b", "neighbouring"),
    (r"\bflavor\b", "flavour"), (r"\bflavors\b", "flavours"), (r"\bflavored\b", "flavoured"),
    (r"\bhumor\b", "humour"), (r"\brumor\b", "rumour"), (r"\blabor\b", "labour"), (r"\bvapor\b", "vapour"),
    (r"\borganize\b", "organise"), (r"\borganizes\b", "organises"), (r"\borganized\b", "organised"), (r"\borganizing\b", "organising"),
    (r"\brecognize\b", "recognise"), (r"\brecognizes\b", "recognises"), (r"\brecognized\b", "recognised"), (r"\brecognizing\b", "recognising"),
    (r"\bapologize\b", "apologise"), (r"\bapologizes\b", "apologises"), (r"\bapologized\b", "apologised"), (r"\bapologizing\b", "apologising"),
    (r"\bcenter\b", "centre"), (r"\bcenters\b", "centres"), (r"\bcentered\b", "centred"), (r"\bcentering\b", "centring"),
    (r"\btheater\b", "theatre"), (r"\btheaters\b", "theatres"),
]

def normalize_en_nz(text: str) -> str:
    """Convert American English spellings to NZ English in display text.

    Applied to all LLM output and STT transcripts before they reach the UI.
    The inverse of normalize_tts_nz(): this converts spoken forms (e.g., "one one one")
    back to digits ("111") for written display, and fixes US spellings.
    """
    # Defensive type handling — Ollama streams can occasionally return non-string types
    if text is None:
        return ""
    if isinstance(text, bytes):
        try:
            text = text.decode("utf-8", errors="replace")
        except Exception:
            text = text.decode("latin-1", errors="replace")
    elif not isinstance(text, str):
        text = str(text)

    # The TTS layer converts "111" → "one one one" for pronunciation.
    # This reverses that for on-screen display and Word document exports.
    try:
        text = re.sub(r"\btriple\s+one\b", "111", text, flags=re.IGNORECASE)
        text = re.sub(r"\bone\s+one\s+one\b", "111", text, flags=re.IGNORECASE)
    except TypeError:
        text = str(text)
        text = re.sub(r"\btriple\s+one\b", "111", text, flags=re.IGNORECASE)
        text = re.sub(r"\bone\s+one\s+one\b", "111", text, flags=re.IGNORECASE)
    def _repl(m: re.Match, target: str) -> str:
        return _preserve_case(m.group(0), target)
    out = text
    for pat, rep in _NZ_SPELLING_RULES:
        try:
            out = re.sub(pat, lambda m, r=rep: _repl(m, r), out, flags=re.IGNORECASE)
        except TypeError:
            # Coerce and retry once if an unexpected type creeps in from streaming
            out = str(out) if out is not None else ""
            out = re.sub(pat, lambda m, r=rep: _repl(m, r), out, flags=re.IGNORECASE)
    return out

# --- Unicode/Mojibake repair & punctuation normaliser (server-side safety net) ---
def normalize_unicode_safe(text: str) -> str:
    """Repair common UTF‑8→Windows‑1252 mojibake and normalise punctuation.
    Returns a plain-UTF‑8-friendly string (mostly ASCII punctuation) so both
    on-screen HTML and python-docx stay clean regardless of upstream encoding.
    """
    if text is None:
        return ""
    if isinstance(text, bytes):
        try:
            text = text.decode("utf-8", errors="replace")
        except Exception:
            text = text.decode("latin-1", errors="replace")
    else:
        text = str(text)

    # 1) Fix typical mojibake sequences (UTF‑8 mis-decoded as cp1252)
    repl = (
        ("â€™", "'"), ("â€˜", "'"), ("â€œ", '"'), ("â€\x9d", '"'), ("â€", '"'),
        ("â€“", "-"), ("â€”", "--"), ("â€¦", "..."), ("Â ", " ")
    )
    for a,b in repl:
        text = text.replace(a,b)

    # 2) Replace actual Unicode punctuation with ASCII-safe forms
    # Curly quotes/apostrophes
    text = re.sub(r"[\u2018\u2019\u2032]", "'", text)
    text = re.sub(r"[\u201C\u201D\u2033]", '"', text)
    # Dashes, ellipsis, NBSP
    text = text.replace("\u2014", "--").replace("\u2013", "-")
    text = text.replace("\u2026", "...")
    text = text.replace("\u00A0", " ")

    # 3) Repair replacement character (U+FFFD) heuristically
    text = re.sub(r"([A-Za-z0-9])\uFFFD([A-Za-z0-9])", r"\1'\2", text)
    text = re.sub(r"\b\uFFFD([A-Za-z])", r"'\1", text)
    text = re.sub(r"([A-Za-z])\uFFFD\b", r"\1'", text)
    text = text.replace("\uFFFD", "'")

    return text

import json
from io import BytesIO
from datetime import datetime
import requests

from fastapi import Request, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document

from backend.models.llm import llm_handler
import threading
import time
from spellchecker import SpellChecker

# Directory containing Piper ONNX voice models and their JSON config files.
# Override with the PIPER_VOICES_DIR environment variable if voices are stored elsewhere.
BASE_DIR = os.path.dirname(__file__)
DEFAULT_VOICES_DIR = os.path.join(BASE_DIR, "models", "audio", "tts", "voices")
VOICES_DIR = os.getenv("PIPER_VOICES_DIR", DEFAULT_VOICES_DIR)

# Base English spell checker; NZ-specific and proper nouns are loaded from custom_dictionary.txt below.
spell = SpellChecker(language='en')

# Path to the admin-editable custom word list for the spell checker.
CUSTOM_DICT_PATH = os.path.join(BASE_DIR, "custom_dictionary.txt")

def load_custom_dictionary():
    """Load words from custom_dictionary.txt and add them to the spell checker.

    The custom dictionary supplements the base English dictionary with NZ-specific
    words, proper nouns (place names, Māori words), and domain terms that would
    otherwise be flagged as misspellings during persona validation.
    Called at startup and again when the admin updates the dictionary via the API.
    """
    custom_words = []
    if os.path.exists(CUSTOM_DICT_PATH):
        try:
            with open(CUSTOM_DICT_PATH, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        custom_words.append(line.lower())
            spell.word_frequency.load_words(custom_words)
            print(f"[API] ✅ Loaded {len(custom_words)} words from custom dictionary")
        except Exception as e:
            print(f"[API] ⚠️ Failed to load custom dictionary: {e}")
    else:
        print(f"[API] ⚠️ Custom dictionary not found at {CUSTOM_DICT_PATH}")
    return custom_words

load_custom_dictionary()

# Load suspicious words list (words that should always trigger warnings)
SUSPICIOUS_WORDS_PATH = os.path.join(BASE_DIR, "suspicious_words.txt")
suspicious_words_set = set()

def load_suspicious_words():
    """Load words from suspicious_words.txt and store them in the global set.

    Suspicious words are always flagged during persona validation regardless of
    whether the spell checker considers them correctly spelled — for example,
    common placeholder values like "lorem" or scenario terms that should not
    appear in a finished persona file.
    Called at startup and when the admin updates the list via the API.
    """
    global suspicious_words_set
    suspicious_words_set = set()
    if os.path.exists(SUSPICIOUS_WORDS_PATH):
        try:
            with open(SUSPICIOUS_WORDS_PATH, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        word = line.split('#')[0].strip().lower()
                        if word:
                            suspicious_words_set.add(word)
            print(f"[API] ✅ Loaded {len(suspicious_words_set)} suspicious words")
        except Exception as e:
            print(f"[API] ⚠️ Failed to load suspicious words: {e}")
    else:
        print(f"[API] ⚠️ Suspicious words file not found at {SUSPICIOUS_WORDS_PATH}")
    return suspicious_words_set

# Load suspicious words on startup
load_suspicious_words()

def check_persona_spelling(persona_data):
    """Check spelling across key persona text fields and return a list of warnings.

    Runs during persona validation (/api/validate-persona) to help instructors catch
    typos before an interview session. Never blocks the upload — warnings are advisory.
    Returns a list of dicts: [{field, word, suggestions}, ...]
    """
    spelling_warnings = []

    try:
        # Fields to check
        fields_to_check = {
            'full_name': persona_data.get('full_name', ''),
            'home_address': persona_data.get('home_address', ''),
            'business_address': persona_data.get('business_address', ''),
            'occupation': persona_data.get('occupation', ''),
            'employed_by': persona_data.get('employed_by', ''),
            'interview_instructions': persona_data.get('interview_instructions', ''),
            'persona_prompt': persona_data.get('persona_prompt', '')
        }

        # Check facts_to_provide array
        facts = persona_data.get('facts_to_provide', [])
        for i, fact_item in enumerate(facts):
            if isinstance(fact_item, dict):
                fact_text = fact_item.get('fact', '')
                reason_text = fact_item.get('reason', '')
                if fact_text and fact_text.strip().lower() not in ['', 'none']:
                    fields_to_check[f'fact_{i+1}'] = fact_text
                if reason_text and reason_text.strip().lower() not in ['', 'none']:
                    fields_to_check[f'fact_{i+1}_reason'] = reason_text

        # Check each field
        for field_name, text in fields_to_check.items():
            if not text or not isinstance(text, str) or text.strip().lower() in ['', 'none']:
                continue

            # The negative lookbehind avoids matching the tail of contractions — e.g.,
            # the "'t" in "don't" would otherwise be extracted as a standalone word.
            # Curly apostrophes (U+2019) are included because LLM output often uses them.
            words = re.findall(r"(?<!['\u2019])\b[a-zA-Z]+(?:['\u2019][a-zA-Z]+)*\b", text)

            for word in words:
                # pyspellchecker only recognises straight apostrophes, so normalise before lookup.
                word_normalized = word.replace('\u2019', "'")
                word_lower = word_normalized.lower()

                # Skip common English contractions (I've, don't, can't, etc.)
                # But allow brand names like Pak'nSave
                if "'" in word_normalized:
                    # Common contraction patterns to skip (already normalized to straight apostrophes)
                    common_contractions = [
                        "'ve", "'re", "'ll", "'d", "'m", "'s", "n't"
                    ]
                    # Check if this is a contraction we should skip
                    is_contraction = any(word_lower.endswith(ending) for ending in common_contractions)
                    if is_contraction:
                        continue
                    # If it has an apostrophe but isn't a common contraction,
                    # strip apostrophes and check the base word (Pak'nSave → paknsave)
                    word_no_apostrophe = word_normalized.replace("'", "").lower()
                    if word_no_apostrophe in spell:
                        continue  # Base word is in dictionary
                    # Continue checking this word without apostrophes
                    word_lower = word_no_apostrophe

                # Skip if it's all caps (likely an acronym)
                if word.isupper() and len(word) > 1:
                    continue

                # Skip single letters
                if len(word) < 2:
                    continue

                # Skip if it's a number-like pattern
                if re.search(r'\d', word):
                    continue

                # Check if word is in suspicious words list (always flag these)
                if word_lower in suspicious_words_set:
                    # Skip proper nouns (capitalized words longer than 3 chars)
                    is_proper_noun = word[0].isupper() and len(word) > 3
                    if not is_proper_noun:
                        spelling_warnings.append({
                            'field': field_name.replace('_', ' ').title(),
                            'word': word,
                            'suggestions': []  # No suggestions for suspicious words
                        })
                    continue

                # Check if misspelled
                if word_lower in spell:
                    continue  # Word is correctly spelled

                # Get suggestions
                misspelled_set = spell.unknown([word_lower])
                if word_lower in misspelled_set:
                    # Get candidates (might be None, so handle that)
                    candidates = spell.candidates(word_lower)
                    suggestions = list(candidates)[:3] if candidates else []

                    # Skip proper nouns (capitalized words longer than 3 chars, like "Priya", "Wellington")
                    # But still report lowercase typos even without suggestions
                    is_proper_noun = word[0].isupper() and len(word) > 3

                    if not is_proper_noun:
                        spelling_warnings.append({
                            'field': field_name.replace('_', ' ').title(),
                            'word': word,
                            'suggestions': suggestions
                        })

    except Exception as e:
        # Spell checking is advisory — a failure must never block persona upload.
        print(f"[API] ⚠️ Spell check error: {e}")
        import traceback
        traceback.print_exc()
        return []

    return spelling_warnings

llm_ready = False

def warmup_llm_on_background():
    """Pre-load the Ollama model into memory to reduce latency on the first interview question.

    Triggered manually via POST /backend/system-warmup-trigger — never called automatically
    at startup, because loading the model can take 10–30 seconds and would slow the app boot.
    Runs in a daemon thread so it does not block the API.
    """
    global llm_ready
    try:
        print("[API] 🔧 LLM Warm-up (background thread) started")
        # A GET to /api/tags forces Ollama to acknowledge the model without generating any text.
        try:
            requests.get("http://127.0.0.1:11434/api/tags", timeout=2)
        except Exception as e:
            print(f"[API] (Non-fatal) Ollama nudge failed: {e}")
        llm_ready = True
        print("[API] ✅ LLM Warm-up complete")
    except Exception as e:
        print(f"[API] ❌ LLM Warm-up failed: {e}")
        llm_ready = False


router = APIRouter()


@router.get("/api/health")
def health_check():
    return {"status": "ok"}

# --- Voices API Endpoint ---
from backend.models.audio.audio_handler import load_voice_list

@router.get("/api/voices")
def get_voice_list():
    voice_data = load_voice_list()
    for entry in voice_data:
        if "filename" not in entry or not entry["filename"]:
            print(f"[WARNING] Missing filename in voice entry: {entry}")
    return voice_data


# --- Curated Piper Voices API Endpoint ---
import csv

@router.get("/api/piper-voices")
def get_piper_voice_list():
    csv_path = os.path.join(VOICES_DIR, "piper_curated_audio.csv")
    # (No change in filename; VOICES_DIR now points to models/tts/voices by default)
    if not os.path.isdir(VOICES_DIR):
        print(f"[API] ⚠️ Voices directory not found at {VOICES_DIR}. Set PIPER_VOICES_DIR to override.")
    voices = []
    try:
        with open(csv_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                voices.append(row)
    except Exception as e:
        print(f"[API] ❌ Failed to read curated Piper voice list: {e}")
        raise HTTPException(status_code=500, detail="Unable to load voice list.")
    return voices



# --- Voice Playback Endpoint for Voice Demo Page (Piper-based) ---
from fastapi.responses import FileResponse

class VoicePreviewRequest(BaseModel):
    text: str
    voice_file: str
    speaker_id: int | None = None

@router.post("/api/play-voice")
async def play_voice_preview(request: VoicePreviewRequest):
    # Validate input
    if not request.voice_file.endswith(".onnx"):
        raise HTTPException(status_code=400, detail="Voice file must be a .onnx file")

    onnx_path = os.path.join(VOICES_DIR, request.voice_file)
    json_path = onnx_path + ".json"

    if not os.path.isfile(onnx_path):
        raise HTTPException(status_code=404, detail=f"ONNX not found: {onnx_path}")
    if not os.path.isfile(json_path):
        raise HTTPException(status_code=404, detail=f"Config JSON not found: {json_path}")

    try:
        preview_text = normalize_tts_nz(request.text)
        wav_bytes = synthesize_speech(
            text=preview_text,
            onnx_path=onnx_path,
            config_path=json_path,
            speaker_id=request.speaker_id
        )
        buffer = BytesIO(wav_bytes)
        buffer.seek(0)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"TTS preview failed: {e}")

    return StreamingResponse(buffer, media_type="audio/wav", headers={"Cache-Control": "no-store"})


# ————————————————————————————————————————————
# ✅ TRANSCRIPT WORD DOCUMENT GENERATION Pt1
# ————————————————————————————————————————————

class TranscriptItem(BaseModel):
    question: str
    answer: str

class TranscriptRequest(BaseModel):
    transcript: list[TranscriptItem]

@router.post("/api/export-transcript")
def export_transcript(request: TranscriptRequest):
    doc = Document()
    doc.add_heading("Witness Interview Transcript", level=1)

    for entry in request.transcript:
        q = normalize_unicode_safe(normalize_en_nz(entry.question))
        a = normalize_unicode_safe(normalize_en_nz(entry.answer))
        doc.add_paragraph(q)
        doc.add_paragraph(a)
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"transcript_{timestamp}.docx"

    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={filename}"})

# —————————————————————————————————
# ✅ PERSONA FILE VALIDATION LOGIC
# —————————————————————————————————

@router.post("/api/validate-persona")
async def validate_persona(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        uploaded_data = json.loads(contents.decode("utf-8"))
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to read file")

    # Stage 1: Key set must exactly match the persona template (no extra or missing keys).
    template_path = os.path.join(os.path.dirname(__file__), "persona_template.json")
    try:
        with open(template_path, "r") as f:
            template_data = json.load(f)
    except Exception:
        raise HTTPException(status_code=500, detail="Unable to load template")

    template_keys = set(template_data.keys())
    uploaded_keys = set(uploaded_data.keys())

    if uploaded_keys != template_keys:
        raise HTTPException(status_code=400, detail="Missing required fields")

    # Stage 2: Reject if every field is blank or "none" (e.g., an unedited template copy).
    non_empty_fields = []
    for key, value in uploaded_data.items():
        if isinstance(value, str):
            if value.strip().lower() != "" and value.strip().lower() != "none":
                non_empty_fields.append(key)
        elif isinstance(value, list):
            if any(
                (isinstance(v, str) and v.strip().lower() not in ["", "none"]) or
                (isinstance(v, dict) and any(
                    vv.strip() != "" and vv.strip().lower() != "none"
                    for vv in v.values() if isinstance(vv, str)
                ))
                for v in value
            ):
                non_empty_fields.append(key)

    if not non_empty_fields:
        raise HTTPException(status_code=400, detail="All fields empty")

    # Stage 3: The eight key fields are required for a working interview — reject if any are blank.
    key_fields = [
        "persona_type",
        "persona_voice_model",
        "persona_voice_speaker_id",
        "full_name",
        "date_of_birth",
        "home_address",
        "interview_instructions",
        "persona_prompt"
    ]
    empty_keys = []
    for k in key_fields:
        v = uploaded_data.get(k, None)
        if v is None:
            empty_keys.append(k)
        elif isinstance(v, str) and v.strip() == "":
            empty_keys.append(k)
    if empty_keys:
        raise HTTPException(status_code=400, detail="Key field empty")

    # Stage 4: Warn (but still allow) if more than half of the optional fields are blank.
    # A persona with few optional fields will work but may produce thin interviews.
    optional_fields = list(set(template_keys) - set(key_fields))
    empty_optional_count = 0

    for key in optional_fields:
        value = uploaded_data.get(key, "")
        if isinstance(value, str):
            if value.strip().lower() in ["", "none"]:
                empty_optional_count += 1
        elif isinstance(value, list):
            if all(
                (isinstance(v, str) and v.strip().lower() in ["", "none"]) or
                (isinstance(v, dict) and all(
                    val.strip() == "" or val.strip().lower() == "none"
                    for val in v.values() if isinstance(val, str)
                ))
                for v in value
            ):
                empty_optional_count += 1

    spelling_warnings = check_persona_spelling(uploaded_data)

    if empty_optional_count > len(optional_fields) / 2:
        uploaded_data["warning"] = True
        if spelling_warnings:
            uploaded_data["spelling_warnings"] = spelling_warnings
        return uploaded_data

    uploaded_data["status"] = "valid_persona_file"
    if spelling_warnings:
        uploaded_data["spelling_warnings"] = spelling_warnings
    return uploaded_data
    
# ————————————————————————————————————————————
# ✅ TRANSCRIPT WORD DOCUMENT GENERATION Pt2
# ————————————————————————————————————————————

# Export transcript as a formatted Word document
class TranscriptExportRequest(BaseModel):
    persona_name: str
    full_name: str
    interview_date: str
    transcript_text: str

@router.post("/api/export-transcript-docx")
def export_transcript_docx(request: TranscriptExportRequest):
    doc = Document()
    # Apply Heading 1 style
    doc.add_heading(f"Interview transcript of {request.full_name} conducted on {request.interview_date}", level=1)

    safe_text = normalize_unicode_safe(normalize_en_nz(request.transcript_text))
    lines = safe_text.strip().split('\n')
    for idx, line in enumerate(lines):
        stripped_line = line.strip()
        line_lower = stripped_line.lower()
        if (
            line_lower.startswith("interview commenced")
            or line_lower.startswith("interview ended")
            or line_lower.startswith("interview recommenced")
            or line_lower.startswith("interview concluded")
        ):
            # Timestamp lines (e.g., "Interview commenced at 14:32") are bolded and unindented.
            # A blank line is added after "commenced"/"recommenced" to separate from the dialogue.
            para = doc.add_paragraph()
            para.paragraph_format.space_after = 0
            para.paragraph_format.left_indent = None
            run = para.add_run(stripped_line)
            run.bold = True
            if "commenced" in line_lower or "recommenced" in line_lower:
                doc.add_paragraph()
        elif line_lower.startswith("[officer]"):
            # Officer speech: bold, full width (no indent), no spacing after
            content = stripped_line[len("[OFFICER]"):].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_after = 0
            para.paragraph_format.left_indent = None
            run = para.add_run(content)
            run.bold = True
        elif line_lower.startswith("[witness]") or line_lower.startswith("[suspect]"):
            # Interviewee speech: indented, normal weight, followed by a blank line for readability
            content = stripped_line.split("]", 1)[-1].strip()
            para = doc.add_paragraph()
            para.paragraph_format.space_after = 6
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run(content)
            run.bold = False
            doc.add_paragraph()
        else:
            # Any other line (e.g., free-form notes) is written as-is
            para = doc.add_paragraph()
            para.paragraph_format.space_after = 6
            run = para.add_run(stripped_line)
            run.bold = False

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    interview_date_str = request.interview_date.strip().replace("_", "-").replace("/", "-")
    parsed_date = datetime.strptime(interview_date_str, "%Y-%m-%d")
    safe_date = parsed_date.strftime("%Y-%m-%d")
    safe_persona_name = request.persona_name.replace(" ", "-")
    filename = f"Interview-{safe_persona_name}-{safe_date}.docx"

    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})

# ————————————————————————————————————————————
# ✅ CUSTOM DICTIONARY MANAGEMENT ENDPOINTS
# ————————————————————————————————————————————

class DictionaryUpdateRequest(BaseModel):
    content: str

@router.get("/api/suspicious-words")
def get_suspicious_words():
    """Get the contents of the suspicious words file."""
    if not os.path.exists(SUSPICIOUS_WORDS_PATH):
        return {"content": "", "exists": False}

    try:
        with open(SUSPICIOUS_WORDS_PATH, 'r', encoding='utf-8') as f:
            content = f.read()
        return {"content": content, "exists": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read suspicious words: {e}")

@router.post("/api/suspicious-words")
def update_suspicious_words(request: DictionaryUpdateRequest):
    """Update the suspicious words file and reload it."""
    try:
        # Write the new content to the file
        with open(SUSPICIOUS_WORDS_PATH, 'w', encoding='utf-8') as f:
            f.write(request.content)

        # Reload the suspicious words
        words = load_suspicious_words()

        return {"status": "success", "words_loaded": len(words)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update suspicious words: {e}")

@router.get("/api/custom-dictionary")
def get_custom_dictionary():
    """Get the contents of the custom dictionary file."""
    if not os.path.exists(CUSTOM_DICT_PATH):
        return {"content": "", "exists": False}

    try:
        with open(CUSTOM_DICT_PATH, 'r', encoding='utf-8') as f:
            content = f.read()
        return {"content": content, "exists": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read custom dictionary: {e}")

@router.post("/api/custom-dictionary")
def update_custom_dictionary(request: DictionaryUpdateRequest):
    """Update the custom dictionary file and reload it."""
    try:
        # Write the new content to the file
        with open(CUSTOM_DICT_PATH, 'w', encoding='utf-8') as f:
            f.write(request.content)

        # Reload the custom dictionary into the spell checker
        # First, we need to reinitialize to clear old custom words
        global spell
        spell = SpellChecker(language='en')
        custom_words = load_custom_dictionary()  # Load the updated custom words

        return {"status": "success", "words_loaded": len(custom_words)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update custom dictionary: {e}")

# Chat API models and route
class ChatRequest(BaseModel):
    message_history: list[dict]
    persona: dict
    audio_mode: bool = False


@router.post("/api/chat")
async def chat_with_persona(request: ChatRequest):
    try:
        response = llm_handler.ask_ollama(request.message_history, request.persona)
        nz_response = normalize_en_nz(response)
        nz_response = normalize_unicode_safe(nz_response)
        result = {"response": nz_response}

        if request.audio_mode:
            voices_dir = VOICES_DIR
            voice_filename = (request.persona or {}).get("persona_voice_model", "")
            if not voice_filename:
                raise HTTPException(status_code=400, detail="Persona is missing 'persona_voice_model'.")

            onnx_path = os.path.join(voices_dir, voice_filename)
            json_path = onnx_path + ".json"

            if not (os.path.isfile(onnx_path) and os.path.isfile(json_path)):
                raise HTTPException(
                    status_code=404,
                    detail=f"🔴 The selected voice model couldn’t be found on this system: {voice_filename}. Please install the voice or choose another."
                )

            try:
                speaker_id_value = request.persona.get("persona_voice_speaker_id")
                try:
                    speaker_id_cast = int(speaker_id_value) if speaker_id_value is not None and str(speaker_id_value).strip() != "" else None
                except Exception:
                    speaker_id_cast = None

                tts_text = normalize_tts_nz(nz_response)
                wav_bytes = synthesize_speech(
                    text=tts_text,
                    onnx_path=onnx_path,
                    config_path=json_path,
                    speaker_id=speaker_id_cast
                )
                buffer = BytesIO(wav_bytes)
                buffer.seek(0)
            except Exception as e:
                import traceback
                traceback.print_exc()
                raise HTTPException(status_code=500, detail=f"TTS generation failed: {e}")

            # HTTP headers are limited to latin-1, so we cannot send NZ text (macrons, etc.) directly.
            # The text is Base64-encoded and placed in a custom header; the frontend decodes it.
            reply_text = nz_response
            reply_text_b64 = base64.b64encode(reply_text.encode("utf-8")).decode("ascii")

            headers = {
                "Cache-Control": "no-store",
                "X-Reply-Text-Base64": reply_text_b64,
                "Access-Control-Expose-Headers": "X-Reply-Text-Base64"
            }
            return StreamingResponse(buffer, media_type="audio/wav", headers=headers)

        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Audio transcription endpoint
from fastapi import Form

@router.post("/api/transcribe-audio")
async def transcribe_audio_api(file: UploadFile = File(...)):
    print("🔹 /api/transcribe-audio endpoint was triggered")
    try:
        from backend.models.audio.audio_handler import transcribe_audio_file
        transcription = transcribe_audio_file(file)
        transcription = normalize_en_nz(transcription)
        transcription = normalize_unicode_safe(transcription)

        # Whisper occasionally hallucinates trailing quote characters at the end of utterances.
        # These look plausible but are artefacts — strip them before returning the transcript.
        transcription = re.sub(r'["\'\u201c\u201d\u2018\u2019]$', '', transcription.strip())
        transcription = re.sub(r'(\?)[\"\'\u201c\u201d\u2018\u2019]$', r'\1', transcription)

        return {"text": transcription}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Transcription failed: {e}")

@router.get("/backend/system-warmup")
def system_warmup_status():
    return {"ready": llm_ready}

# —————————————————
# ✅ SYSTEM CHECK
# —————————————————
import platform
import subprocess


@router.get("/backend/system-check")
def system_check():
    import importlib
    import subprocess
    import traceback

    def get_installed_version(pkg_name):
        try:
            import importlib.metadata as metadata
            return metadata.version(pkg_name)
        except importlib.metadata.PackageNotFoundError:
            return "N/A"

    def get_version_by_import(module_name):
        return get_installed_version(module_name)

    def get_version_by_command(command):
        try:
            output = subprocess.check_output(command, shell=True, stderr=subprocess.STDOUT, text=True)
            if command.strip().endswith("pip --version"):
                parts = output.strip().split()
                for part in parts:
                    if part[0].isdigit():
                        return part
                return "N/A"
            if "python" in command:
                return output.strip().split()[-1]
            if "ollama" in command:
                return output.strip().split()[-1]
            if "ffmpeg" in command:
                # Extract version number from: ffmpeg version 7.1.1 ...
                first_line = output.strip().splitlines()[0]
                parts = first_line.split()
                if len(parts) >= 3 and parts[0].lower() == "ffmpeg" and parts[1].lower() == "version":
                    return parts[2]  # e.g., '7.1.1'
                return "N/A"
            return output.strip()
        except Exception:
            return "N/A"

    # --- Faster-Whisper small.en model inspection (local + remote) ---
    def _read_json(path):
        try:
            import json as _json
            with open(path, "r", encoding="utf-8") as f:
                return _json.load(f)
        except Exception:
            return None

    def _read_text_40(path):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()[:40]
        except Exception:
            return None

    def get_small_en_model_status():
        """
        Return a dict describing the local Faster-Whisper small.en model *weights* and
        whether they match the remote 'main' head. This is separate from the Python package.
        Prefers a marker '.hf_snapshot.json' (new layout), then falls back to old refs/snapshots.
        """
        base_dir = os.path.dirname(__file__)
        model_dir = os.path.join(base_dir, "models", "audio", "stt", "whisper",
                                 "models--Systran--faster-whisper-small.en")

        exists = os.path.isdir(model_dir)
        local_commit = None

        # 1) Preferred: marker written by our installer (flat layout)
        marker_path = os.path.join(model_dir, ".hf_snapshot.json")
        marker = _read_json(marker_path)
        if isinstance(marker, dict):
            local_commit = marker.get("commit_sha") or marker.get("sha")

        # 2) Legacy HF cache-style layout
        if not local_commit:
            ref_path = os.path.join(model_dir, "refs", "main")
            ref_hash = _read_text_40(ref_path)
            if ref_hash and len(ref_hash) == 40:
                local_commit = ref_hash

        # 3) Last resort: try to infer a commit by reading the single folder in snapshots/
        if not local_commit:
            try:
                snaps_dir = os.path.join(model_dir, "snapshots")
                if os.path.isdir(snaps_dir):
                    entries = [d for d in os.listdir(snaps_dir)
                               if len(d) == 40 and os.path.isdir(os.path.join(snaps_dir, d))]
                    if entries:
                        local_commit = entries[0]
            except Exception:
                pass

        # Remote HEAD on Hugging Face (best-effort; may be offline)
        remote_commit = None
        try:
            from huggingface_hub import HfApi
            api = HfApi()
            info = api.model_info("Systran/faster-whisper-small.en", revision="main")
            remote_commit = getattr(info, "sha", None)
        except Exception:
            remote_commit = None

        up_to_date = (local_commit == remote_commit) if (local_commit and remote_commit) else None

        return {
            "path_exists": exists,
            "local_commit": local_commit or "unknown",
            "remote_head": remote_commit or "unknown",
            "up_to_date": up_to_date
        }

    tools = {
        "python": platform.python_version(),
        "fastapi": get_installed_version("fastapi"),
        "uvicorn": get_installed_version("uvicorn"),
        "pydantic": get_installed_version("pydantic"),
        "ollama": get_version_by_command("ollama --version"),
        "torch": get_installed_version("torch"),
        "ctranslate2": get_installed_version("ctranslate2"),
        "faster-whisper": get_installed_version("faster-whisper"),
        "piper-tts": get_installed_version("piper-tts"),
        "sounddevice": get_installed_version("sounddevice"),
        "python-docx": get_installed_version("python-docx"),
        "python-multipart": get_installed_version("python-multipart"),
        "pip": get_version_by_command(f"{sys.executable} -m pip --version"),
        "pydub": get_installed_version("pydub"),
        "ffmpeg": get_version_by_command("ffmpeg -version"),
    }

    try:
        response = {}
        for name, version in tools.items():
            # Handle non-PyPI tools
            if name == "ffmpeg":
                # ffmpeg is not on PyPI; attempt to query the installed package manager for the latest version.
                try:
                    latest_ffmpeg_version = None
                    try:
                        # Try brew (macOS); returns non-zero on other platforms, handled gracefully below
                        brew_proc = subprocess.run(
                            ["brew", "info", "--json=v2", "ffmpeg"],
                            capture_output=True, text=True, timeout=5
                        )
                        if brew_proc.returncode == 0:
                            info_json = brew_proc.stdout
                            import json as json_lib
                            info = json_lib.loads(info_json)
                            # Find latest version from brew info
                            latest_formula = info.get("formulae", [{}])[0]
                            latest_ffmpeg_version = latest_formula.get("versions", {}).get("stable", "unknown")
                        else:
                            latest_ffmpeg_version = "unknown"
                    except Exception:
                        latest_ffmpeg_version = "unknown"
                    latest_version = latest_ffmpeg_version
                except Exception:
                    latest_version = "unknown"
            elif name == "ollama":
                try:
                    output = subprocess.check_output(
                        "curl -s https://api.github.com/repos/ollama/ollama/releases/latest",
                        shell=True, text=True, timeout=5
                    )
                    import json as json_lib
                    latest_version = json_lib.loads(output).get("tag_name", "Unknown").lstrip("v")
                except Exception:
                    latest_version = "Unknown"
            elif name == "python":
                try:
                    output = subprocess.check_output(
                        "curl -s https://endoflife.date/api/python.json",
                        shell=True, text=True, timeout=5
                    )
                    import json as json_lib
                    versions = json_lib.loads(output)
                    version_nums = [v.get("latest", "") for v in versions if v.get("latest")]
                    if version_nums:
                        from packaging.version import parse as parse_version
                        latest_version = str(max(map(parse_version, version_nums)))
                    else:
                        latest_version = "Unknown"
                except Exception:
                    latest_version = "Unknown"
            else:
                try:
                    latest_output = subprocess.check_output(
                        f"pip index versions {name}", shell=True,
                        stderr=subprocess.DEVNULL, text=True, timeout=5
                    )
                    latest_version = "Unknown"
                    for line in latest_output.splitlines():
                        if "AVAILABLE VERSIONS" in line.upper():
                            parts = line.split(":")
                            if len(parts) > 1:
                                version_candidates = [v.strip() for v in parts[1].split(",") if v.strip()]
                                if version_candidates:
                                    latest_version = version_candidates[0]
                            break
                except Exception:
                    latest_version = "Unknown"

            response[name] = {
                "status": "ok" if version not in ["N/A", "unknown"] else "error",
                "installed_version": version,
                "latest_version": latest_version,
                "version_state": "unknown" if name == "ffmpeg" else (
                    "latest" if latest_version == version else
                    "outdated" if latest_version not in ["Unknown", ""] else "unknown"
                )
            }


        # Add a separate entry for the Faster-Whisper small.en model weights (distinct from the Python package version).
        small_en_status = get_small_en_model_status()
        response["faster-whisper-model-small.en"] = {
            "status": "ok" if small_en_status.get("path_exists") else "error",
            "installed_version": small_en_status.get("local_commit", "unknown"),
            "latest_version": small_en_status.get("remote_head", "unknown"),
            "version_state": (
                "latest" if small_en_status.get("up_to_date") is True else
                "outdated" if small_en_status.get("up_to_date") is False else
                "unknown"
            )
        }

        return response

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="System check failed")


@router.get("/backend/persona-template")
def get_persona_template():
    template_path = os.path.join(os.path.dirname(__file__), "persona_template.json")
    try:
        with open(template_path, "r", encoding="utf-8") as f:
            template_data = json.load(f)
        return template_data
    except Exception:
        raise HTTPException(status_code=500, detail="Unable to load persona template")

# Register the router with the FastAPI app
from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
import os

app = FastAPI()

@app.post("/backend/system-warmup-trigger")
def trigger_warmup():
    thread = threading.Thread(target=warmup_llm_on_background, daemon=True)
    thread.start()
    return {"status": "llm_warmup_started"}

app.include_router(router)

# Serve the entire frontend directory as static files.
# The "/" mount must come LAST — FastAPI evaluates mounts in registration order,
# and a wildcard static mount would shadow any API routes registered after it.
app.mount(
    "/",
    StaticFiles(directory=os.path.join(os.path.dirname(__file__), "..", "frontend"), html=True),
    name="static"
)
# Helper function to clean up old audio files
import glob

def clean_old_audio_files(folder="frontend/temp_audio", age_limit_seconds=300):
    """Remove TTS audio files older than 5 minutes (300 seconds)"""
    now = time.time()
    for filepath in glob.glob(os.path.join(folder, "response_*.wav")):
        if os.path.isfile(filepath):
            file_age = now - os.path.getmtime(filepath)
            if file_age > age_limit_seconds:
                try:
                    os.remove(filepath)
                except Exception as e:
                    print(f"[Cleanup] Failed to delete old TTS file {filepath}: {e}")